# 1) https://pypi.org/project/streamlit/

# 2) Terminal: cd PYTHON/EcxelToPdf > pip install streamlit pandas python-docx matplotlib openpyxl
#                                   > streamlit run app.py

# 3) Welcome to Streamlit! / You can now view your Streamlit app in your browser.
        # Local URL: http://localhost:8501
        # Network URL: http://100.126.10.200:8501

#-------------------------------------------------------------------------------
#-------------------------------------------------------------------------------

import streamlit as st
import pandas as pd                 # Biblioteca para trabajar con datos en formato Excel (y CSV)
from docx import Document           # Biblioteca para crear y manipular documentos de Word
from docx.shared import Inches      # Para especificar dimensiones en pulgadas dentro del documento Word
import matplotlib.pyplot as plt     # Biblioteca para crear gráficos
import io                           # Para trabajar con flujos de entrada/salida de datos en memoria

#-----------------------------------------
# Función principal para crear el informe
#-----------------------------------------
def create_report(template_path, data, chart_data=None):
    st.write("Iniciando la creación del informe...")  # Mensaje informativo en la interfaz de Streamlit
    doc = Document(template_path)                     # Abrir el documento Word de plantilla

    #-------------------------------------------------------------
    # Reemplazar marcadores de posición en el texto del documento
    #-------------------------------------------------------------
    for paragraph in doc.paragraphs:                # Iterar sobre cada párrafo del documento
        for key, value in data.items():             # Iterar sobre cada par clave-valor de los datos proporcionados
            if f'{{{{{key}}}}}' in paragraph.text:  # Buscar marcadores de posición en el formato '{{clave}}'
                st.write(f"Reemplazando {key} con {value} en el informe.")  # Mensaje informativo
                paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))  # Reemplazar el marcador con el valor correspondiente

    #-------------------------------------------------------------
    # Generar e insertar el gráfico si se proporcionaron datos para ello
    #-------------------------------------------------------------
    if chart_data is not None:
        st.write("Generando gráfico...")    # Mensaje informativo
        plt.figure(figsize=(6, 4))          # Crear una figura para el gráfico con tamaño de 6x4 pulgadas
        plt.bar(chart_data['labels'], chart_data['values'])  # Crear un gráfico de barras con las etiquetas y valores proporcionados
        plt.title(chart_data['title'])      # Asignar un título al gráfico
        plt.xlabel(chart_data.get('xlabel', ''))  # Asignar una etiqueta al eje X (si se proporcionó)
        plt.ylabel(chart_data.get('ylabel', ''))  # Asignar una etiqueta al eje Y (si se proporcionó)
        img_buffer = io.BytesIO()           # Crear un búfer en memoria para guardar la imagen del gráfico

        plt.savefig(img_buffer, format='png')  # Guardar el gráfico como imagen PNG en el búfer
        img_buffer.seek(0)  # Mover el puntero del búfer al inicio para poder leerlo desde el principio

        st.write("Insertando gráfico en el marcador del documento...")  # Mensaje informativo
        for paragraph in doc.paragraphs:  # Iterar sobre los párrafos para buscar el marcador del gráfico
            if '[Gráfico aqui]' in paragraph.text:  # Buscar el marcador '[Gráfico aqui]'
                paragraph.text = paragraph.text.replace('[Gráfico aqui]', '')   # Eliminar el marcador
                paragraph.add_run().add_picture(img_buffer, width=Inches(6))    # Insertar la imagen del gráfico con un ancho de 6 pulgadas

    output = io.BytesIO()   # Crear un búfer en memoria para guardar el informe generado
    doc.save(output)        # Guardar el documento modificado en el búfer
    output.seek(0)          # Mover el puntero del búfer al inicio
    st.write("Informe creado con exito")  # Mensaje informativo
    return output           # Devolver el búfer con el informe generado

def main():
    #--------------------------------
    # Configuración de la página
    #--------------------------------
    st.set_page_config(
        page_title="Generador de Informes",   # Título de la página
        page_icon="images/EC-IT-LOGO-1.png",  # Ícono de la página
        layout="centered"  # Centrar el contenido de la página
    )
    st.subheader("Generador de Informes")  # Subtítulo en la página

    #-------------------------------------------------------------
    # Widgets para cargar los archivos de plantilla y datos
    #-------------------------------------------------------------
    template_file = st.file_uploader("Cargar plantilla Word", type="docx")  # Solo permite cargar archivos .docx
    data_file = st.file_uploader("Cargar datos", type=["xlsx", "csv"])      # Permite cargar archivos Excel o CSV

    #-------------------------------------------------------------
    # Verificar si se cargaron ambos archivos
    #-------------------------------------------------------------
    if template_file and data_file:
        st.success("Archivos cargados correctamente")  # Mensaje de éxito

        # Leer el archivo de datos según su tipo
        if data_file.name.endswith('.csv'):
            df = pd.read_csv(data_file)    # Leer archivo CSV con pandas
        else:
            df = pd.read_excel(data_file)  # Leer archivo Excel con pandas

        st.subheader("Datos cargados")
        st.dataframe(df)  # Mostrar los datos en una tabla en la interfaz

        #-------------------------------------------------------------
        # Selector de fila para elegir los datos a usar en el informe
        #-------------------------------------------------------------
        row_index = st.selectbox("Seleccionar fila para el informe", options=range(len(df)))
        selected_data = df.iloc[row_index].to_dict()  # Obtener los datos de la fila seleccionada como un diccionario

        #-------------------------------------------------------------
        # Checkbox para habilitar la generación del gráfico
        #-------------------------------------------------------------
        generate_chart = st.checkbox("Generar gráfico")
        chart_data = None  # Inicializar variable para los datos del gráfico

        #-------------------------------------------------------------
        # Si se marca la opción de generar gráfico, obtener los datos necesarios
        #-------------------------------------------------------------
        if generate_chart:
            chart_title = st.text_input("Título del gráfico", "Gráfico de Datos")  # Título del gráfico (con valor predeterminado)
            x_column = st.selectbox("Columna para eje X", options=df.columns)  # Selector para elegir la columna del eje X
            y_column = st.selectbox("Columna para eje Y", options=df.columns)  # Selector para elegir la columna del eje Y

            #-------------------------------------------------------------
            # Crear un diccionario con los datos del gráfico
            #-------------------------------------------------------------
            chart_data = {
                'title': chart_title,
                'labels': df[x_column].tolist(),  # Valores de la columna seleccionada para el eje X
                'values': df[y_column].tolist(),  # Valores de la columna seleccionada para el eje Y
                'xlabel': x_column,  # Etiqueta del eje X
                'ylabel': y_column   # Etiqueta del eje Y
            }

        #------------------------------
        # Botón para generar el informe
        #------------------------------
        if st.button("Generar Informe"):
            output = create_report(template_file, selected_data, chart_data)  # Llamar a la función para crear el informe
            # Botón para descargar el informe generado
            st.download_button(
                "Descargar informe",  # Texto del botón
                output,  # Datos del informe
                "informe_generado.docx",  # Nombre del archivo descargado
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"  # Tipo MIME para Word
            )

    #-------------------------------------------------------------
    # Si falta alguno de los archivos, mostrar un mensaje de advertencia
    #-------------------------------------------------------------
    elif template_file or data_file:
        st.warning("Por favor, carga ambos archivos: la plantilla Word y Ecxel.")

#-------------------------------------------------------------
# Ejecutar la función principal al iniciar la aplicación
#-------------------------------------------------------------
main()

